"""
Document Extraction Service for Project RAG

Extracts text content from various document formats:
- PDF
- DOCX
- TXT
- CSV

IMPORTANT: This service is part of the RAG pipeline which is DISABLED by default.
RAG features require larger models (≥14B/70B) to work safely in healthcare context.
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any
import PyPDF2
import csv
import io

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """
    Extracts text content from uploaded documents.
    Supports PDF, DOCX, TXT, and CSV files.
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.csv', '.xlsx'}
    
    def __init__(self, project_files_dir: str = "project_files"):
        self.project_files_dir = Path(project_files_dir)
        self.project_files_dir.mkdir(parents=True, exist_ok=True)
    
    def extract_text(self, file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            file_type: Optional file type override (pdf, docx, txt, csv)
            
        Returns:
            Dict with:
                - success: bool
                - text: extracted text content
                - metadata: file metadata
                - error: error message if failed
        """
        path = Path(file_path)
        
        if not path.exists():
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": f"File not found: {file_path}"
            }
        
        ext = file_type or path.suffix.lower()
        
        try:
            if ext == '.pdf' or ext == 'pdf':
                return self._extract_pdf(path)
            elif ext == '.docx' or ext == 'docx':
                return self._extract_docx(path)
            elif ext == '.txt' or ext == 'txt':
                return self._extract_txt(path)
            elif ext == '.csv' or ext == 'csv':
                return self._extract_csv(path)
            elif ext == '.xlsx' or ext == 'xlsx':
                return self._extract_xlsx(path)
            else:
                return {
                    "success": False,
                    "text": "",
                    "metadata": {"extension": ext},
                    "error": f"Unsupported file type: {ext}"
                }
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ext},
                "error": str(e)
            }
    
    def _extract_pdf(self, path: Path) -> Dict[str, Any]:
        """Extract text from PDF file."""
        text_parts = []
        num_pages = 0
        
        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            full_text = "\n\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "metadata": {
                    "extension": ".pdf",
                    "num_pages": num_pages,
                    "char_count": len(full_text),
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ".pdf"},
                "error": f"PDF extraction failed: {str(e)}"
            }
    
    def _extract_docx(self, path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            # Try to import python-docx
            try:
                from docx import Document
            except ImportError:
                return {
                    "success": False,
                    "text": "",
                    "metadata": {"extension": ".docx"},
                    "error": "python-docx not installed. Run: pip install python-docx"
                }
            
            doc = Document(path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "metadata": {
                    "extension": ".docx",
                    "num_paragraphs": len(doc.paragraphs),
                    "char_count": len(full_text),
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ".docx"},
                "error": f"DOCX extraction failed: {str(e)}"
            }
    
    def _extract_txt(self, path: Path) -> Dict[str, Any]:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                return {
                    "success": False,
                    "text": "",
                    "metadata": {"extension": ".txt"},
                    "error": "Could not decode file with any supported encoding"
                }
            
            return {
                "success": True,
                "text": text,
                "metadata": {
                    "extension": ".txt",
                    "char_count": len(text),
                    "line_count": text.count('\n') + 1,
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ".txt"},
                "error": f"TXT extraction failed: {str(e)}"
            }
    
    def _extract_csv(self, path: Path) -> Dict[str, Any]:
        """Extract text from CSV file."""
        try:
            text_parts = []
            row_count = 0
            
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f)
                for row in reader:
                    row_text = " | ".join(str(cell).strip() for cell in row if str(cell).strip())
                    if row_text:
                        text_parts.append(row_text)
                        row_count += 1
            
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "metadata": {
                    "extension": ".csv",
                    "row_count": row_count,
                    "char_count": len(full_text),
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except Exception as e:
            logger.error(f"CSV extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ".csv"},
                "error": f"CSV extraction failed: {str(e)}"
            }
    
    def _extract_xlsx(self, path: Path) -> Dict[str, Any]:
        """Extract text from XLSX file."""
        try:
            try:
                import openpyxl
            except ImportError:
                return {
                    "success": False,
                    "text": "",
                    "metadata": {"extension": ".xlsx"},
                    "error": "openpyxl not installed. Run: pip install openpyxl"
                }
            
            wb = openpyxl.load_workbook(path, data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                
                for row in sheet.iter_rows():
                    row_text = " | ".join(
                        str(cell.value).strip() 
                        for cell in row 
                        if cell.value is not None and str(cell.value).strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "metadata": {
                    "extension": ".xlsx",
                    "sheet_count": len(wb.sheetnames),
                    "char_count": len(full_text),
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except Exception as e:
            logger.error(f"XLSX extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ".xlsx"},
                "error": f"XLSX extraction failed: {str(e)}"
            }
    
    def save_extracted_text(self, project_id: str, doc_id: str, text: str) -> str:
        """
        Save extracted text to the project documents directory.
        
        Args:
            project_id: Project ID
            doc_id: Document ID
            text: Extracted text content
            
        Returns:
            Path to the saved text file
        """
        project_dir = self.project_files_dir / project_id
        project_dir.mkdir(parents=True, exist_ok=True)
        
        text_file = project_dir / f"{doc_id}.txt"
        with open(text_file, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return str(text_file)
