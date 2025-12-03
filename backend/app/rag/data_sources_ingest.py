"""
Data Source Ingestion Pipeline

Handles extraction and chunking of content from various file formats:
- PDF, DOCX, TXT, MD (text extraction)
- CSV (row extraction)
- JSON (field extraction)

Each chunk is tagged with metadata for category-based retrieval.
"""

import os
import json
import csv
import io
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import uuid

# Import file extraction dependencies with graceful fallbacks
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from app.config import settings

logger = logging.getLogger(__name__)


# Category keywords for automatic detection
CATEGORY_KEYWORDS = {
    "oncology": [
        "cancer", "tumor", "oncology", "chemotherapy", "carcinoma", "metastasis",
        "radiation", "immunotherapy", "leukemia", "lymphoma", "melanoma", "sarcoma",
        "biopsy", "malignant", "benign", "neoplasm"
    ],
    "cardiology": [
        "heart", "cardiac", "cardiology", "cardiovascular", "arrhythmia", "hypertension",
        "myocardial", "infarction", "angina", "stroke", "atherosclerosis", "cholesterol",
        "blood pressure", "ecg", "echocardiogram", "pacemaker"
    ],
    "diabetes": [
        "diabetes", "insulin", "glucose", "glycemic", "hba1c", "metformin", "glp-1",
        "hyperglycemia", "hypoglycemia", "pancreas", "islet", "type 2 diabetes",
        "type 1 diabetes", "diabetic", "blood sugar"
    ],
    "neurology": [
        "neurology", "brain", "neurological", "alzheimer", "parkinson", "epilepsy",
        "seizure", "dementia", "multiple sclerosis", "neuropathy", "migraine",
        "stroke", "cognitive", "neural", "neuron"
    ],
    "immunology": [
        "immune", "immunology", "autoimmune", "inflammation", "antibody", "antigen",
        "lymphocyte", "cytokine", "allergy", "immunotherapy", "vaccine", "t-cell",
        "b-cell", "immunodeficiency"
    ],
    "respiratory": [
        "lung", "respiratory", "pulmonary", "asthma", "copd", "pneumonia", "bronchitis",
        "emphysema", "tuberculosis", "oxygen", "ventilator", "inhaler", "breathing"
    ],
    "gastroenterology": [
        "gastro", "liver", "hepatic", "intestine", "colon", "ibd", "crohn", "colitis",
        "cirrhosis", "hepatitis", "pancreatic", "digestive", "stomach", "gi tract"
    ],
    "internal_knowledge": [
        "internal", "proprietary", "company", "protocol", "sop", "guideline",
        "procedure", "policy", "training", "documentation"
    ],
    "clinical_trials": [
        "clinical trial", "phase 1", "phase 2", "phase 3", "phase 4", "rct",
        "randomized", "placebo", "endpoint", "efficacy", "safety", "enrollment",
        "protocol", "investigator", "fda", "ema"
    ],
    "pharmacology": [
        "pharmacology", "drug", "medication", "dosage", "pharmacokinetics", "pkpd",
        "absorption", "metabolism", "excretion", "half-life", "bioavailability",
        "drug interaction", "adverse event", "side effect"
    ]
}


class DataSourceIngester:
    """
    Ingests files from data sources, extracts text, and creates tagged chunks.
    """
    
    def __init__(self):
        self.chunk_size = settings.DATA_SOURCE_CHUNK_SIZE
        self.chunk_overlap = settings.DATA_SOURCE_CHUNK_OVERLAP
        self.data_dir = Path(settings.DATA_SOURCES_DIR)
        self.data_dir.mkdir(parents=True, exist_ok=True)
        
    def detect_category(self, text: str, explicit_category: Optional[str] = None) -> str:
        """
        Detect the category of content based on keywords.
        If explicit_category is provided, use that instead.
        """
        if explicit_category:
            return explicit_category.lower().replace(" ", "_")
        
        text_lower = text.lower()
        category_scores = {}
        
        for category, keywords in CATEGORY_KEYWORDS.items():
            score = sum(1 for kw in keywords if kw in text_lower)
            if score > 0:
                category_scores[category] = score
        
        if category_scores:
            return max(category_scores, key=category_scores.get)
        
        return "general"
    
    def extract_text(self, file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract text content from a file.
        
        Returns:
            Dict with:
                - success: bool
                - text: extracted text
                - metadata: file metadata
                - rows: (for CSV) list of row dicts
                - fields: (for JSON) extracted fields
                - error: error message if failed
        """
        path = Path(file_path)
        
        if not path.exists():
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": f"File not found: {file_path}"
            }
        
        ext = file_type or path.suffix.lower()
        
        try:
            if ext in ['.pdf', 'pdf']:
                return self._extract_pdf(path)
            elif ext in ['.docx', 'docx']:
                return self._extract_docx(path)
            elif ext in ['.txt', 'txt', '.md', 'md']:
                return self._extract_txt(path)
            elif ext in ['.csv', 'csv']:
                return self._extract_csv(path)
            elif ext in ['.json', 'json']:
                return self._extract_json(path)
            else:
                return {
                    "success": False,
                    "text": "",
                    "metadata": {"extension": ext},
                    "error": f"Unsupported file type: {ext}"
                }
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ext},
                "error": str(e)
            }
    
    def _extract_pdf(self, path: Path) -> Dict[str, Any]:
        """Extract text from PDF file."""
        if not HAS_PYPDF2:
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ".pdf"},
                "error": "PyPDF2 not installed. Run: pip install PyPDF2"
            }
        
        text_parts = []
        num_pages = 0
        
        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text.strip())
            
            full_text = "\n\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "metadata": {
                    "extension": ".pdf",
                    "num_pages": num_pages,
                    "char_count": len(full_text),
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ".pdf"},
                "error": f"PDF extraction failed: {str(e)}"
            }
    
    def _extract_docx(self, path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        if not HAS_DOCX:
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ".docx"},
                "error": "python-docx not installed. Run: pip install python-docx"
            }
        
        try:
            doc = Document(path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "metadata": {
                    "extension": ".docx",
                    "num_paragraphs": len(doc.paragraphs),
                    "char_count": len(full_text),
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": ".docx"},
                "error": f"DOCX extraction failed: {str(e)}"
            }
    
    def _extract_txt(self, path: Path) -> Dict[str, Any]:
        """Extract text from TXT/MD file."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return {
                "success": True,
                "text": text,
                "metadata": {
                    "extension": path.suffix.lower(),
                    "char_count": len(text),
                    "line_count": text.count('\n') + 1,
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return {
                    "success": True,
                    "text": text,
                    "metadata": {
                        "extension": path.suffix.lower(),
                        "char_count": len(text),
                        "encoding": "latin-1",
                        "file_size": path.stat().st_size
                    },
                    "error": None
                }
            except Exception as e:
                return {
                    "success": False,
                    "text": "",
                    "metadata": {"extension": path.suffix.lower()},
                    "error": f"Text extraction failed: {str(e)}"
                }
        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "metadata": {"extension": path.suffix.lower()},
                "error": f"Text extraction failed: {str(e)}"
            }
    
    def _extract_csv(self, path: Path) -> Dict[str, Any]:
        """Extract rows from CSV file."""
        try:
            rows = []
            text_parts = []
            
            with open(path, 'r', encoding='utf-8', newline='') as f:
                reader = csv.DictReader(f)
                headers = reader.fieldnames or []
                
                for row in reader:
                    rows.append(row)
                    # Create text representation of row
                    row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "rows": rows,
                "metadata": {
                    "extension": ".csv",
                    "row_count": len(rows),
                    "columns": headers,
                    "char_count": len(full_text),
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except Exception as e:
            logger.error(f"CSV extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "rows": [],
                "metadata": {"extension": ".csv"},
                "error": f"CSV extraction failed: {str(e)}"
            }
    
    def _extract_json(self, path: Path) -> Dict[str, Any]:
        """Extract fields from JSON file."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Flatten JSON to text
            text_parts = []
            
            def flatten_json(obj, prefix=""):
                if isinstance(obj, dict):
                    for k, v in obj.items():
                        new_prefix = f"{prefix}.{k}" if prefix else k
                        flatten_json(v, new_prefix)
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        flatten_json(item, f"{prefix}[{i}]")
                else:
                    text_parts.append(f"{prefix}: {obj}")
            
            flatten_json(data)
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "fields": data,
                "metadata": {
                    "extension": ".json",
                    "field_count": len(text_parts),
                    "char_count": len(full_text),
                    "file_size": path.stat().st_size
                },
                "error": None
            }
        except Exception as e:
            logger.error(f"JSON extraction error: {str(e)}")
            return {
                "success": False,
                "text": "",
                "fields": {},
                "metadata": {"extension": ".json"},
                "error": f"JSON extraction failed: {str(e)}"
            }
    
    def create_chunks(
        self,
        text: str,
        source_id: str,
        source_name: str,
        file_id: str,
        category: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Split text into chunks with metadata tags.
        
        Args:
            text: The full extracted text
            source_id: ID of the data source
            source_name: Name of the data source
            file_id: ID of the specific file
            category: Optional explicit category (auto-detected if not provided)
            metadata: Additional metadata to attach to each chunk
        
        Returns:
            List of chunk dictionaries with text and metadata
        """
        if not text or not text.strip():
            return []
        
        # Detect category if not explicitly provided
        detected_category = self.detect_category(text, category)
        
        # Clean and normalize text
        text = re.sub(r'\s+', ' ', text).strip()
        
        chunks = []
        chunk_id = 0
        start = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            # If we're not at the end, try to find a good break point
            if end < len(text):
                # Look for sentence end (., !, ?) within the last 200 chars
                break_zone = text[max(start, end - 200):end]
                
                # Find the last sentence-ending punctuation
                last_period = break_zone.rfind('. ')
                last_question = break_zone.rfind('? ')
                last_exclaim = break_zone.rfind('! ')
                
                best_break = max(last_period, last_question, last_exclaim)
                
                if best_break > 0:
                    end = max(start, end - 200) + best_break + 2
                else:
                    # Fall back to last space
                    last_space = text[start:end].rfind(' ')
                    if last_space > 0:
                        end = start + last_space
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk = {
                    "chunk_id": f"{file_id}_chunk_{chunk_id}",
                    "text": chunk_text,
                    "source_id": source_id,
                    "source_name": source_name,
                    "category": detected_category,
                    "file_id": file_id,
                    "upload_date": datetime.utcnow().isoformat(),
                    "char_start": start,
                    "char_end": end,
                    "chunk_index": chunk_id,
                    "metadata": metadata or {}
                }
                chunks.append(chunk)
                chunk_id += 1
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            
            # Prevent infinite loop
            if start >= len(text) - self.chunk_overlap:
                break
        
        logger.info(f"Created {len(chunks)} chunks for file {file_id} in category '{detected_category}'")
        return chunks
    
    async def ingest_file(
        self,
        file_path: str,
        source_id: str,
        source_name: str,
        category: Optional[str] = None,
        file_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Full ingestion pipeline for a single file.
        
        Args:
            file_path: Path to the file
            source_id: ID of the data source
            source_name: Name of the data source
            category: Optional category (auto-detected if not provided)
            file_id: Optional file ID (generated if not provided)
            metadata: Additional metadata
        
        Returns:
            Dict with ingestion results including chunks
        """
        if not file_id:
            file_id = str(uuid.uuid4())
        
        # Extract text
        extraction_result = self.extract_text(file_path)
        
        if not extraction_result["success"]:
            return {
                "success": False,
                "file_id": file_id,
                "chunks": [],
                "error": extraction_result["error"]
            }
        
        # Create chunks
        chunks = self.create_chunks(
            text=extraction_result["text"],
            source_id=source_id,
            source_name=source_name,
            file_id=file_id,
            category=category,
            metadata={
                **(metadata or {}),
                "file_path": file_path,
                "extraction_metadata": extraction_result["metadata"]
            }
        )
        
        return {
            "success": True,
            "file_id": file_id,
            "chunks": chunks,
            "chunk_count": len(chunks),
            "category": chunks[0]["category"] if chunks else (category or "general"),
            "extraction_metadata": extraction_result["metadata"],
            "error": None
        }


# Singleton instance
data_source_ingester = DataSourceIngester()
